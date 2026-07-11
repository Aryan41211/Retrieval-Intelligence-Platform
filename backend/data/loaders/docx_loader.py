"""DOCX document loader."""

import zipfile
from pathlib import Path

from backend.core.exceptions import DocumentLoadError, EmptyDocumentError
from backend.data.loaders.base_loader import BaseLoader
from backend.data.models.document import Document

_DOCX_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _extract_text_from_docx_zip(path: str | Path) -> str:
    """Fallback: extract text directly from a DOCX zip archive.

    Used when python-docx fails due to lxml compatibility issues.
    """
    import re

    from lxml import etree

    NS = _DOCX_NS

    paragraphs: list[str] = []
    with zipfile.ZipFile(str(path), "r") as zf:
        if "word/document.xml" not in zf.namelist():
            raise DocumentLoadError(f"Invalid DOCX: missing word/document.xml in {path}")
        raw = zf.read("word/document.xml")
        # Inject the namespace declaration when the XML omits it.
        if b"<w:" in raw and b"xmlns:w" not in raw:
            raw = re.sub(
                rb"(<\w+:document\b)",
                rb'\1 xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
                raw,
                count=1,
            )
        root = etree.fromstring(raw)
        for para in root.iter(f"{{{NS}}}p"):
            texts: list[str] = []
            for t in para.iter(f"{{{NS}}}t"):
                if t.text:
                    texts.append(t.text)
            line = "".join(texts).strip()
            if line:
                paragraphs.append(line)
    return "\n\n".join(paragraphs)


class DOCXLoader(BaseLoader):
    """Load DOCX (Word) documents."""

    def get_supported_extensions(self) -> list[str]:
        """Return list of supported extensions.

        Returns:
            List of supported file extensions.
        """
        return [".docx"]

    def load(self, file_path: str | Path) -> Document:
        """Load a DOCX document.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Document object with extracted text.

        Raises:
            DocumentLoadError: If the document cannot be loaded.
            EmptyDocumentError: If the document has no text content.
        """
        path = self.validate_file(file_path)
        checksum = self.compute_checksum(path)
        file_extension = path.suffix.lower()

        if path.stat().st_size == 0:
            raise DocumentLoadError(f"File is empty: {path}")

        try:
            import docx  # python-docx
        except ImportError:
            raise DocumentLoadError(
                "python-docx not installed. Install with: pip install python-docx"
            ) from None

        try:
            doc = docx.Document(str(path))
        except AttributeError as e:
            if "overrides" in str(e):
                # python-docx / lxml compatibility issue — fall back to
                # manual extraction from the zip archive.
                content = _extract_text_from_docx_zip(path)
                if not content.strip():
                    raise EmptyDocumentError(f"DOCX contains no text content: {path}") from e
                return self._build_document(path, content, checksum, file_extension)
            raise DocumentLoadError(f"Failed to read DOCX: {e}") from e
        except Exception as e:
            raise DocumentLoadError(f"Failed to read DOCX: {e}") from e

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        if not content.strip():
            raise EmptyDocumentError(f"DOCX contains no text content: {path}")

        title = None
        author = None
        created_at = None

        if doc.core_properties:
            title = doc.core_properties.title or None
            author = doc.core_properties.author or None
            if doc.core_properties.created:
                created_at = doc.core_properties.created

        page_count = None

        metadata = self.build_metadata(
            title=title,
            author=author,
            char_count=len(content),
            word_count=len(content.split()),
            page_count=page_count,
        )

        source = self.build_source(path, checksum)

        document = Document(
            filename=path.name,
            file_extension=file_extension,
            source_path=str(path),
            content=content,
            metadata=metadata,
            checksum=checksum,
            file_size=path.stat().st_size,
            source=source,
        )

        if created_at:
            document = document.model_copy(update={"created_at": created_at})

        return document

    def _build_document(
        self,
        path: Path,
        content: str,
        checksum: str,
        file_extension: str,
    ) -> Document:
        """Build a Document from extracted content (fallback path)."""
        metadata = self.build_metadata(
            char_count=len(content),
            word_count=len(content.split()),
        )
        source = self.build_source(path, checksum)
        return Document(
            filename=path.name,
            file_extension=file_extension,
            source_path=str(path),
            content=content,
            metadata=metadata,
            checksum=checksum,
            file_size=path.stat().st_size,
            source=source,
        )
